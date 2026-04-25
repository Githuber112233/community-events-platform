# -*- coding: utf-8 -*-
"""
生成本科毕业论文中期检查评价表 Word 文档
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUTPUT = r"D:\desktop\工作实习\毕设\本科毕业论文中期检查评价表_柳鸿博.docx"

doc = Document()

# ── 页面边距 ──────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.left_margin   = Cm(3)
section.right_margin  = Cm(2.5)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ── 辅助函数 ──────────────────────────────────────────────
def set_font(run, name="宋体", size=12, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_para(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT,
             font_name="宋体", font_size=12, bold=False, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    set_font(run, font_name, font_size, bold)
    return p

def set_cell_bg(cell, hex_color="D9D9D9"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def cell_text(cell, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
              vertical=WD_ALIGN_VERTICAL.CENTER, font="宋体", color=None):
    cell.vertical_alignment = vertical
    p = cell.paragraphs[0]
    p.clear()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, font, size, bold, color)

def add_cell_para(cell, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, font="宋体"):
    """在单元格中追加一个段落"""
    p = cell.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    set_font(run, font, size, bold)
    return p

def merge_vertical(table, col, start_row, end_row):
    table.cell(start_row, col).merge(table.cell(end_row, col))

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

# ══════════════════════════════════════════════════════════
#  标题
# ══════════════════════════════════════════════════════════
add_para(doc, "西北农林科技大学", align=WD_ALIGN_PARAGRAPH.CENTER,
         font_name="黑体", font_size=16, bold=True, space_before=6, space_after=4)
add_para(doc, "信息工程学院 2026 届本科生毕业论文", align=WD_ALIGN_PARAGRAPH.CENTER,
         font_name="黑体", font_size=14, bold=True, space_before=0, space_after=4)
add_para(doc, "中期检查评价表", align=WD_ALIGN_PARAGRAPH.CENTER,
         font_name="黑体", font_size=16, bold=True, space_before=0, space_after=10)

# ══════════════════════════════════════════════════════════
#  基本信息表（5列）
# ══════════════════════════════════════════════════════════
tbl1 = doc.add_table(rows=4, cols=6)
tbl1.style = "Table Grid"
tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER

# 列宽
widths = [2.5, 3.5, 2.5, 3.5, 2.0, 3.0]
for r in tbl1.rows:
    for i, w in enumerate(widths):
        r.cells[i].width = Cm(w)

def info_row(row, pairs):
    """pairs: [(label, value), ...]，每对占2列"""
    for i, (label, value) in enumerate(pairs):
        cell_text(row.cells[i*2],   label, bold=True, size=10.5, font="黑体")
        cell_text(row.cells[i*2+1], value, size=10.5)
        set_cell_bg(row.cells[i*2], "E8E8E8")

info_row(tbl1.rows[0], [("论文题目", "社区兴趣活动平台设计与实现"), ("专业班级", "计算机2202班"), ("学号", "2022013268")])
info_row(tbl1.rows[1], [("学生姓名", "柳鸿博"), ("指导教师", "王莉"), ("职称", "中级职称")])
info_row(tbl1.rows[2], [("所属院系", "信息工程学院"), ("检查日期", "2026年4月初"), ("检查方式", "导师面谈+系统演示")])
# 合并最后一行做备注
tbl1.rows[3].cells[0].merge(tbl1.rows[3].cells[5])
cell_text(tbl1.rows[3].cells[0],
          "说明：本表由指导教师在论文中期检查时填写，对学生的阶段进展进行综合评价。",
          size=9.5, font="宋体")

doc.add_paragraph()  # 空行

# ══════════════════════════════════════════════════════════
#  一、工作进度与计划执行情况
# ══════════════════════════════════════════════════════════
add_para(doc, "一、工作进度与计划执行情况", font_name="黑体", font_size=12, bold=True,
         space_before=6, space_after=4)

tbl2 = doc.add_table(rows=9, cols=5)
tbl2.style = "Table Grid"
tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER

col_ws = [0.8, 3.0, 3.2, 3.2, 2.5]
for r in tbl2.rows:
    for i, w in enumerate(col_ws):
        r.cells[i].width = Cm(w)

# 表头
headers = ["序号", "计划阶段", "计划完成时间", "实际完成情况", "完成状态"]
for i, h in enumerate(headers):
    cell_text(tbl2.rows[0].cells[i], h, bold=True, size=10.5,
              align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体")
    set_cell_bg(tbl2.rows[0].cells[i], "BFBFBF")

# 数据行
rows_data = [
    ("1", "选题与开题准备",        "2025.12.09–2026.01.07", "已完成文献调研、需求分析，开题报告获批通过",           "✅ 已完成"),
    ("2", "系统需求分析与总体设计", "2026.01.09–2026.02.05", "已完成系统架构设计、数据库设计（6张核心表）、接口规范文档", "✅ 已完成"),
    ("3", "数据库与后端基础功能实现","2026.02.06–2026.03.05", "已实现用户模块（JWT登录/注册）、活动管理、评论点赞等后端接口，Redis缓存集成完毕", "✅ 已完成"),
    ("4", "前端功能实现与系统联调",  "2026.03.06–2026.04.05", "已完成前端核心页面（首页、活动列表、活动详情、创建活动、个人中心），前后端联调基本完成", "🔄 进行中（基本完成）"),
    ("5", "混合推荐算法设计与实现",  "2026.04.06–2026.04.30", "算法方案已完成设计，代码实现阶段即将启动",             "⏳ 待开始"),
    ("6", "系统测试与性能优化",     "2026.05.01–2026.05.15", "未开始",                                           "⏳ 未开始"),
    ("7", "论文撰写与修改",         "2026.05.16–2026.06.05", "已整理部分章节提纲及需求分析初稿",                   "⏳ 准备中"),
]

for idx, (no, stage, time, actual, status) in enumerate(rows_data, 1):
    row = tbl2.rows[idx]
    cell_text(row.cells[0], no,     size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row.cells[1], stage,  size=10.5)
    cell_text(row.cells[2], time,   size=10)
    cell_text(row.cells[3], actual, size=10)
    cell_text(row.cells[4], status, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  二、已完成的主要工作内容
# ══════════════════════════════════════════════════════════
add_para(doc, "二、已完成的主要工作内容", font_name="黑体", font_size=12, bold=True,
         space_before=6, space_after=4)

tbl3 = doc.add_table(rows=6, cols=2)
tbl3.style = "Table Grid"
tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl3.columns[0].width = Cm(3.2)
tbl3.columns[1].width = Cm(13.3)
for row in tbl3.rows:
    row.cells[0].width = Cm(3.2)
    row.cells[1].width = Cm(13.3)

# 表头
cell_text(tbl3.rows[0].cells[0], "工作模块", bold=True, size=10.5,
          align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体")
cell_text(tbl3.rows[0].cells[1], "具体完成情况", bold=True, size=10.5,
          align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体")
set_cell_bg(tbl3.rows[0].cells[0], "BFBFBF")
set_cell_bg(tbl3.rows[0].cells[1], "BFBFBF")

work_items = [
    ("系统架构\n与数据库设计",
     "采用前后端分离架构：前端 Vue 3 + TypeScript + Vite + Tailwind CSS，后端 Spring Boot 3 + JPA + JWT + Redis，数据库 MySQL 8。"
     "设计并建立6张核心数据表：user（用户）、activity（活动）、activity_registration（报名记录）、"
     "activity_like（点赞）、comment（评论）、category（分类），表结构完整，关联关系合理。"
     "\n\n【截图提示①】建议截取数据库表结构截图（如 Navicat/DataGrip 的 ER 图或表列表界面）"),

    ("后端功能\n模块实现",
     "完成以下 RESTful API 接口开发：\n"
     "· 用户模块：注册/登录（基于JWT鉴权）、获取/更新个人信息\n"
     "· 活动管理模块：活动发布（含分类、时间、地点）、活动列表查询、活动详情查询、活动搜索\n"
     "· 社交互动模块：活动评论（发表/查看）、点赞/取消点赞、活动报名/取消报名\n"
     "· 推荐基础：用户行为数据采集接口（浏览记录、报名记录写入）\n"
     "· Redis 缓存热点活动列表，减少数据库压力\n"
     "\n【截图提示②】建议截取 Postman/Apifox 的接口测试成功截图（如登录接口、活动发布接口返回200状态）"),

    ("前端页面\n实现",
     "完成全部核心前端页面开发（Vue 3 单页应用）：\n"
     "· 首页（Home）：Hero Banner、热门活动推荐、分类导航\n"
     "· 活动列表页（Events）：分类筛选、搜索、活动卡片网格展示\n"
     "· 活动详情页（EventDetail）：活动信息、报名按钮、评论区、点赞功能\n"
     "· 创建活动页（CreateEvent）：多步骤表单（基础信息、时间地点、分类标签）\n"
     "· 个人中心页（Profile）：已报名活动、发布的活动管理\n"
     "· 登录/注册页（Login / Register）：表单验证、JWT存储\n"
     "\n【截图提示③】建议截取系统运行截图，包括：首页、活动列表页、活动详情页各一张（浏览器全屏截图）"),

    ("前后端\n联调",
     "完成前后端接口对接，主要联调内容：\n"
     "· 用户鉴权流程（登录获取Token → 前端存储 → 请求携带Authorization头）\n"
     "· 活动数据的增删查（创建活动、首页列表渲染、详情页渲染均已联调通过）\n"
     "· 评论、点赞功能的实时状态更新\n"
     "· CORS跨域问题解决，统一API响应格式（code/data/message）\n"
     "\n【截图提示④】建议截取浏览器开发者工具中Network面板，展示某个API请求的请求头（含Authorization）及响应数据"),

    ("推荐算法\n前期调研",
     "完成混合推荐算法的技术选型与方案设计：\n"
     "· 调研并分析了基于内容的推荐（Content-Based Filtering）与基于用户的协同过滤（User-CF）两种算法的原理与适用场景\n"
     "· 设计了动态权重融合策略：新用户阶段以CB算法为主，随行为数据积累逐步提升User-CF权重\n"
     "· 引入时间衰减因子，解决用户兴趣漂移问题\n"
     "· 确定了用户画像构建方案（兴趣标签Tag + 历史行为评分矩阵）\n"
     "· TF-IDF 活动文本特征提取方案已确定，实现工作下阶段启动"),
]

for idx, (module, content) in enumerate(work_items, 1):
    row = tbl3.rows[idx]
    cell_text(row.cells[0], module, size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER,
              vertical=WD_ALIGN_VERTICAL.CENTER, font="宋体")
    set_cell_bg(row.cells[0], "F2F2F2")
    # 内容单元格——逐行写入（保留换行）
    cell = row.cells[1]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    lines = content.split("\n")
    first = True
    for line in lines:
        if first:
            p = cell.paragraphs[0]
            p.clear()
            first = False
        else:
            p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(line)
        # 截图提示行加粗变色
        if "截图提示" in line:
            set_font(run, "宋体", 10, True, (192, 0, 0))
        else:
            set_font(run, "宋体", 10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  三、论文撰写进度
# ══════════════════════════════════════════════════════════
add_para(doc, "三、论文撰写进度", font_name="黑体", font_size=12, bold=True,
         space_before=6, space_after=4)

tbl4 = doc.add_table(rows=9, cols=4)
tbl4.style = "Table Grid"
tbl4.alignment = WD_TABLE_ALIGNMENT.CENTER
ws4 = [1.0, 4.0, 5.5, 2.0]
for r in tbl4.rows:
    for i, w in enumerate(ws4):
        r.cells[i].width = Cm(w)

chapter_headers = ["章节", "标题", "主要内容", "完成度"]
for i, h in enumerate(chapter_headers):
    cell_text(tbl4.rows[0].cells[i], h, bold=True, size=10.5,
              align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体")
    set_cell_bg(tbl4.rows[0].cells[i], "BFBFBF")

chapters = [
    ("第1章", "绪论",
     "研究背景（社区活动信息碎片化问题）、国内外研究现状（小红书/豆瓣同城/推荐系统文献综述）、研究目标与意义、论文结构安排",
     "约80%"),
    ("第2章", "相关技术综述",
     "B/S架构、Spring Boot、Vue.js、JWT认证、Redis缓存、MySQL数据库、推荐算法基础理论（CB、User-CF）",
     "约60%"),
    ("第3章", "系统需求分析与总体设计",
     "功能需求（用户/组织者/管理员视角）、非功能需求、系统架构图、数据库设计（ER图、表结构）、接口设计规范",
     "约70%"),
    ("第4章", "核心功能模块实现",
     "用户模块、活动管理模块、社交互动模块的详细设计与代码实现，关键代码说明",
     "约40%（待系统完善后补充）"),
    ("第5章", "混合推荐算法设计与实现",
     "算法方案设计、CB算法实现、User-CF实现、动态加权融合策略、实验评估（准确率/召回率等）",
     "约10%（下阶段重点）"),
    ("第6章", "系统测试与性能评估",
     "单元测试、集成测试、JMeter压力测试、测试结果分析",
     "0%（待后续）"),
    ("第7章", "总结与展望",
     "研究总结、创新点归纳、不足与未来工作展望",
     "提纲已列"),
]

for idx, (ch, title, content, pct) in enumerate(chapters, 1):
    row = tbl4.rows[idx]
    cell_text(row.cells[0], ch,      size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(row.cells[1], title,   size=10.5, bold=True)
    cell_text(row.cells[2], content, size=10)
    cell_text(row.cells[3], pct,     size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    if "80%" in pct or "70%" in pct or "60%" in pct:
        set_cell_bg(row.cells[3], "E2EFDA")   # 绿色背景
    elif "40%" in pct or "10%" in pct:
        set_cell_bg(row.cells[3], "FFF2CC")   # 黄色背景
    else:
        set_cell_bg(row.cells[3], "FCE4D6")   # 红色背景

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  四、存在问题与下阶段计划
# ══════════════════════════════════════════════════════════
add_para(doc, "四、存在问题与下阶段工作计划", font_name="黑体", font_size=12, bold=True,
         space_before=6, space_after=4)

tbl5 = doc.add_table(rows=3, cols=2)
tbl5.style = "Table Grid"
tbl5.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in tbl5.rows:
    row.cells[0].width = Cm(3.5)
    row.cells[1].width = Cm(13.0)

# 小标题行1
cell_text(tbl5.rows[0].cells[0], "存在的主要问题", bold=True, size=10.5,
          align=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_VERTICAL.CENTER, font="黑体")
set_cell_bg(tbl5.rows[0].cells[0], "F2F2F2")
problems = tbl5.rows[0].cells[1]
problems.vertical_alignment = WD_ALIGN_VERTICAL.TOP
prob_lines = [
    "（1）推荐算法实现尚未启动，为后续工作的核心难点。混合推荐涉及User-CF中余弦相似度计算的矩阵维度问题，在数据量有限时推荐效果的验证具有一定挑战性。",
    "（2）用户行为数据不足：系统处于开发阶段，缺乏真实用户数据，推荐算法评估需依赖模拟数据集，结果可信度需在论文中说明。",
    "（3）前端个别交互细节（如活动地图选点、二维码签到）尚待完善，不影响核心功能演示。",
    "（4）系统性能压测尚未进行，Redis缓存策略的效果需通过实验数据支撑。",
]
first = True
for line in prob_lines:
    if first:
        p = problems.paragraphs[0]; p.clear(); first = False
    else:
        p = problems.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(line)
    set_font(run, "宋体", 10.5)

# 小标题行2
cell_text(tbl5.rows[1].cells[0], "解决措施", bold=True, size=10.5,
          align=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_VERTICAL.CENTER, font="黑体")
set_cell_bg(tbl5.rows[1].cells[0], "F2F2F2")
solutions = tbl5.rows[1].cells[1]
solutions.vertical_alignment = WD_ALIGN_VERTICAL.TOP
sol_lines = [
    "（1）推荐算法按原定计划于4月上旬启动，先完成CB算法原型，再逐步集成User-CF，保障核心功能按时实现。",
    "（2）使用公开数据集（如MovieLens的评分数据结构）构造模拟的用户-活动评分矩阵用于算法调试与离线评估。",
    "（3）前端次要功能在系统联调稳定后补充完善，优先保证核心演示流程完整。",
    "（4）在系统测试阶段使用JMeter进行压力测试，结合Redis命中率指标优化缓存策略。",
]
first = True
for line in sol_lines:
    if first:
        p = solutions.paragraphs[0]; p.clear(); first = False
    else:
        p = solutions.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(line)
    set_font(run, "宋体", 10.5)

# 小标题行3：下阶段计划
cell_text(tbl5.rows[2].cells[0], "下阶段工作计划\n（4月–5月）", bold=True, size=10.5,
          align=WD_ALIGN_PARAGRAPH.CENTER, vertical=WD_ALIGN_VERTICAL.CENTER, font="黑体")
set_cell_bg(tbl5.rows[2].cells[0], "F2F2F2")
plans = tbl5.rows[2].cells[1]
plans.vertical_alignment = WD_ALIGN_VERTICAL.TOP
plan_lines = [
    "【4月上旬（4.06–4.15）】混合推荐算法实现：完成CB算法（基于兴趣标签TF-IDF匹配）核心代码，集成到后端推荐接口，前端'猜你喜欢'模块联调。",
    "【4月中旬（4.16–4.25）】User-CF算法实现：构建用户-活动评分矩阵，实现余弦相似度计算，完成动态加权融合逻辑，输出推荐列表。",
    "【4月下旬（4.26–4.30）】推荐算法评估：设计离线实验，计算准确率/召回率/F1指标，与基准算法对比，记录实验数据用于论文第五章。",
    "【5月上旬（5.01–5.15）】系统测试与优化：单元测试、集成测试、JMeter压力测试，根据结果优化数据库索引和Redis缓存策略。",
    "【5月中旬起（5.16–）】论文撰写：完成第4、5、6章撰写，全文整合修改，准备答辩材料。",
]
first = True
for line in plan_lines:
    if first:
        p = plans.paragraphs[0]; p.clear(); first = False
    else:
        p = plans.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(line)
    set_font(run, "宋体", 10.5)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  五、截图附件说明
# ══════════════════════════════════════════════════════════
add_para(doc, "五、截图附件说明（需手动补充）", font_name="黑体", font_size=12, bold=True,
         space_before=6, space_after=4)

screenshots_table = doc.add_table(rows=6, cols=3)
screenshots_table.style = "Table Grid"
screenshots_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in screenshots_table.rows:
    row.cells[0].width = Cm(1.0)
    row.cells[1].width = Cm(5.5)
    row.cells[2].width = Cm(10.0)

sc_headers = ["编号", "截图内容", "操作说明"]
for i, h in enumerate(sc_headers):
    cell_text(screenshots_table.rows[0].cells[i], h, bold=True, size=10.5,
              align=WD_ALIGN_PARAGRAPH.CENTER, font="黑体")
    set_cell_bg(screenshots_table.rows[0].cells[i], "BFBFBF")

screenshot_items = [
    ("①", "数据库表结构",
     "打开 Navicat / DataGrip / MySQL Workbench，截取数据库中6张核心表的列表或ER图，\n"
     "确保表名清晰可见（user, activity, activity_registration, activity_like, comment, category）"),
    ("②", "后端接口测试（Postman/Apifox）",
     "打开 Postman 或 Apifox，截取至少2个接口的成功响应截图：\n"
     "· 用户登录接口（POST /api/auth/login）返回JWT token\n"
     "· 活动列表接口（GET /api/activities）返回活动数组\n"
     "截图须包含请求URL、请求体/参数、响应状态码200及响应JSON数据"),
    ("③", "前端系统运行截图",
     "浏览器打开 http://localhost:5173 或 5174，截取以下页面各一张：\n"
     "· 首页（含活动卡片列表）\n"
     "· 活动详情页（含评论区）\n"
     "· 个人中心页\n"
     "使用 Windows 截图工具（Win+Shift+S）或浏览器全页截图"),
    ("④", "前后端联调—网络请求",
     "浏览器F12 → Network → XHR，截取某个API请求详情：\n"
     "· Headers 标签页，展示 Authorization: Bearer xxx 请求头\n"
     "· Preview/Response 标签页，展示服务器返回的JSON数据"),
    ("⑤", "代码量/提交记录（可选加分项）",
     "打开 VS Code 或 IntelliJ IDEA，截取项目文件树；\n"
     "或打开 Git 提交日志（git log --oneline 命令行）展示开发进度；\n"
     "或截取代码编辑器中某个核心代码文件（如 ActivityController.java 或 recommend.js）"),
]

for idx, (no, title, desc) in enumerate(screenshot_items, 1):
    row = screenshots_table.rows[idx]
    cell_text(row.cells[0], no, size=11, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
    cell_text(row.cells[1], title, size=10.5, bold=True, font="黑体")
    set_cell_bg(row.cells[1], "FFF2CC")  # 黄色提示
    # 描述多行
    cell = row.cells[2]
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    lines = desc.split("\n")
    first = True
    for line in lines:
        if first:
            p = cell.paragraphs[0]; p.clear(); first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        run = p.add_run(line)
        set_font(run, "宋体", 10)

doc.add_paragraph()

# ══════════════════════════════════════════════════════════
#  六、指导教师评语与签字
# ══════════════════════════════════════════════════════════
add_para(doc, "六、指导教师评语与综合评定", font_name="黑体", font_size=12, bold=True,
         space_before=6, space_after=4)

tbl6 = doc.add_table(rows=5, cols=4)
tbl6.style = "Table Grid"
tbl6.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in tbl6.rows:
    for i, w in enumerate([3.0, 6.5, 3.0, 4.0]):
        row.cells[i].width = Cm(w)

# 评语区（合并4列）
tbl6.rows[0].cells[0].merge(tbl6.rows[0].cells[3])
cell = tbl6.rows[0].cells[0]
cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
p = cell.paragraphs[0]
p.clear()
run = p.add_run("指导教师评语：")
set_font(run, "黑体", 11, True)
for _ in range(6):
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    p2.paragraph_format.space_after = Pt(2)
    run2 = p2.add_run("　")
    set_font(run2, "宋体", 12)
# 占位提示
p_hint = cell.add_paragraph()
run_hint = p_hint.add_run("（请指导教师根据实际进展在此处填写评语，约200字）")
set_font(run_hint, "宋体", 10, color=(128, 128, 128))

# 评分项
score_items = [
    ("工作态度与学习能力", "□ 优秀  □ 良好  □ 一般  □ 较差"),
    ("按计划完成情况",     "□ 按时完成  □ 基本按时  □ 有所滞后  □ 严重滞后"),
    ("技术方案合理性",     "□ 优秀  □ 良好  □ 一般  □ 较差"),
    ("综合评定",          "□ 优秀  □ 良好  □ 合格  □ 不合格"),
]
for idx, (label, options) in enumerate(score_items, 1):
    tbl6.rows[idx].cells[0].merge(tbl6.rows[idx].cells[1])
    cell_text(tbl6.rows[idx].cells[0], label, bold=True, size=11, font="黑体")
    set_cell_bg(tbl6.rows[idx].cells[0], "F2F2F2")
    tbl6.rows[idx].cells[2].merge(tbl6.rows[idx].cells[3])
    cell_text(tbl6.rows[idx].cells[2], options, size=11)

doc.add_paragraph()

# 签字行
tbl7 = doc.add_table(rows=2, cols=4)
tbl7.style = "Table Grid"
tbl7.alignment = WD_TABLE_ALIGNMENT.CENTER
for row in tbl7.rows:
    for i, w in enumerate([3.5, 4.5, 3.5, 5.0]):
        row.cells[i].width = Cm(w)

cell_text(tbl7.rows[0].cells[0], "指导教师签字：", bold=True, size=11, font="黑体")
cell_text(tbl7.rows[0].cells[1], "（手写签名）", size=11)
cell_text(tbl7.rows[0].cells[2], "检查日期：", bold=True, size=11, font="黑体")
cell_text(tbl7.rows[0].cells[3], "2026年    月    日", size=11)

tbl7.rows[1].cells[0].merge(tbl7.rows[1].cells[3])
cell_text(tbl7.rows[1].cells[0],
          "注：本表一式两份，院系和指导教师各留存一份，作为答辩资格审查依据。",
          size=9.5, font="宋体", color=(128, 128, 128))

# ══════════════════════════════════════════════════════════
#  保存
# ══════════════════════════════════════════════════════════
doc.save(OUTPUT)
print(f"[OK] 文件已保存至: {OUTPUT}")
